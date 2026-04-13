# File: zulong/l1b/document_parser.py
# 文档解析模块 - 将非结构化文件转化为文本或向量
# 对应 TSD 数据解析管道规范

import os
import logging
import asyncio
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass
from concurrent.futures import ThreadPoolExecutor

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    """解析后的文档结构"""
    filepath: str
    filename: str
    file_type: str
    content: str
    chunks: List[str]
    metadata: Dict[str, Any]
    parse_time_ms: float


class DocumentParser:
    """文档解析器
    
    支持格式：
    - 文本文档: .txt, .md, .json, .csv
    - Word 文档: .doc, .docx
    - PDF 文档: .pdf (支持文本和 OCR)
    - 图片: .jpg, .png, .bmp (OCR 或 Caption)
    """
    
    _instance = None
    _executor = ThreadPoolExecutor(max_workers=2)
    
    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
        return cls._instance
    
    def __init__(self):
        if not hasattr(self, '_initialized'):
            self._initialized = False
            self._ocr_engine = None
            self._vision_model = None
            logger.info("[DocumentParser] 初始化完成")
    
    def initialize(self, enable_ocr: bool = True, enable_vision: bool = False):
        """初始化解析器
        
        Args:
            enable_ocr: 是否启用 OCR
            enable_vision: 是否启用视觉模型
        """
        if self._initialized:
            return
        
        if enable_ocr:
            self._init_ocr_engine()
        
        if enable_vision:
            self._init_vision_model()
        
        self._initialized = True
        logger.info(f"[DocumentParser] 初始化完成 (OCR: {self._ocr_engine is not None}, Vision: {self._vision_model is not None})")
    
    def _init_ocr_engine(self):
        """初始化 OCR 引擎"""
        try:
            import fitz
            self._ocr_engine = "pymupdf"
            logger.info("[DocumentParser] OCR 引擎: PyMuPDF")
        except ImportError:
            try:
                import pdfplumber
                self._ocr_engine = "pdfplumber"
                logger.info("[DocumentParser] OCR 引擎: pdfplumber")
            except ImportError:
                try:
                    import pytesseract
                    self._ocr_engine = "tesseract"
                    logger.info("[DocumentParser] OCR 引擎: Tesseract")
                except ImportError:
                    logger.warning("[DocumentParser] 未检测到 OCR 引擎，图片/PDF 解析功能受限")
                    self._ocr_engine = None
    
    def _init_vision_model(self):
        """初始化视觉模型"""
        try:
            from zulong.models.container import ModelContainer
            from zulong.models.config import ModelID
            
            container = ModelContainer()
            self._vision_model = container.resident_models.get(ModelID.VISION)
            
            if self._vision_model:
                logger.info("[DocumentParser] 视觉模型已加载")
            else:
                logger.warning("[DocumentParser] 视觉模型未加载")
        except Exception as e:
            logger.warning(f"[DocumentParser] 视觉模型初始化失败: {e}")
            self._vision_model = None
    
    def parse_file(self, filepath: str, chunk_size: int = 500) -> ParsedDocument:
        """解析文件
        
        Args:
            filepath: 文件路径
            chunk_size: 分块大小
            
        Returns:
            ParsedDocument: 解析后的文档
        """
        import time
        start_time = time.time()
        
        filename = os.path.basename(filepath)
        file_type = self._get_file_type(filepath)
        
        logger.info(f"[DocumentParser] 开始解析: {filename} (类型: {file_type})")
        
        content = self._extract_content(filepath, file_type)
        
        chunks = self._chunk_content(content, chunk_size) if content else []
        
        metadata = self._build_metadata(filepath, file_type, content)
        
        elapsed = (time.time() - start_time) * 1000
        
        logger.info(f"[DocumentParser] 解析完成: {filename}, 内容长度: {len(content)}, 分块数: {len(chunks)}, 耗时: {elapsed:.1f}ms")
        
        return ParsedDocument(
            filepath=filepath,
            filename=filename,
            file_type=file_type,
            content=content,
            chunks=chunks,
            metadata=metadata,
            parse_time_ms=elapsed
        )
    
    async def parse_file_async(self, filepath: str, chunk_size: int = 500) -> ParsedDocument:
        """异步解析文件
        
        Args:
            filepath: 文件路径
            chunk_size: 分块大小
            
        Returns:
            ParsedDocument: 解析后的文档
        """
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            self._executor,
            self.parse_file,
            filepath,
            chunk_size
        )
    
    def _get_file_type(self, filepath: str) -> str:
        """获取文件类型
        
        Args:
            filepath: 文件路径
            
        Returns:
            str: 文件类型
        """
        ext = os.path.splitext(filepath)[1].lower()
        
        type_mapping = {
            '.txt': 'text',
            '.md': 'text',
            '.json': 'json',
            '.csv': 'csv',
            '.doc': 'word',
            '.docx': 'word',
            '.pdf': 'pdf',
            '.jpg': 'image',
            '.jpeg': 'image',
            '.png': 'image',
            '.bmp': 'image',
            '.gif': 'image',
            '.py': 'code',
            '.js': 'code',
            '.html': 'code',
            '.css': 'code'
        }
        
        return type_mapping.get(ext, 'unknown')
    
    def _extract_content(self, filepath: str, file_type: str) -> str:
        """提取文件内容
        
        Args:
            filepath: 文件路径
            file_type: 文件类型
            
        Returns:
            str: 提取的内容
        """
        try:
            if file_type == 'text':
                return self._read_text(filepath)
            elif file_type == 'json':
                return self._read_json(filepath)
            elif file_type == 'csv':
                return self._read_csv(filepath)
            elif file_type == 'word':
                return self._read_word(filepath)
            elif file_type == 'pdf':
                return self._read_pdf(filepath)
            elif file_type == 'image':
                return self._read_image(filepath)
            elif file_type == 'code':
                return self._read_text(filepath)
            else:
                logger.warning(f"[DocumentParser] 不支持的文件类型: {file_type}")
                return ""
        except Exception as e:
            logger.error(f"[DocumentParser] 内容提取失败: {e}")
            return ""
    
    def _read_text(self, filepath: str) -> str:
        """读取文本文件"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
        
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"[DocumentParser] 读取文本失败: {e}")
                return ""
        
        logger.error(f"[DocumentParser] 无法解码文件: {filepath}")
        return ""
    
    def _read_json(self, filepath: str) -> str:
        """读取 JSON 文件"""
        import json
        
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
                return json.dumps(data, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.error(f"[DocumentParser] 读取 JSON 失败: {e}")
            return ""
    
    def _read_csv(self, filepath: str) -> str:
        """读取 CSV 文件"""
        try:
            import pandas as pd
            df = pd.read_csv(filepath)
            return df.to_string()
        except ImportError:
            logger.warning("[DocumentParser] pandas 未安装，使用基础 CSV 解析")
            import csv
            
            lines = []
            with open(filepath, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                for row in reader:
                    lines.append(" | ".join(row))
            return "\n".join(lines)
        except Exception as e:
            logger.error(f"[DocumentParser] 读取 CSV 失败: {e}")
            return ""
    
    def _read_word(self, filepath: str) -> str:
        """读取 Word 文档"""
        try:
            from docx import Document
            doc = Document(filepath)
            
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                paragraphs.append("\n".join(table_text))
            
            return "\n\n".join(paragraphs)
            
        except ImportError:
            logger.error("[DocumentParser] python-docx 未安装")
            return ""
        except Exception as e:
            logger.error(f"[DocumentParser] 读取 Word 失败: {e}")
            return ""
    
    def _read_pdf(self, filepath: str) -> str:
        """读取 PDF 文档
        
        支持文本层和 OCR
        """
        text = ""
        
        try:
            import fitz
            
            doc = fitz.open(filepath)
            
            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                
                if page_text.strip():
                    text += page_text + "\n"
                else:
                    logger.info(f"[DocumentParser] 第 {page_num + 1} 页为扫描件，尝试 OCR")
                    ocr_text = self._ocr_pdf_page(page)
                    text += ocr_text + "\n"
            
            doc.close()
            
        except ImportError:
            logger.warning("[DocumentParser] PyMuPDF 未安装，尝试 pdfplumber")
            
            try:
                import pdfplumber
                
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                            
            except ImportError:
                logger.error("[DocumentParser] PyMuPDF 和 pdfplumber 都未安装")
                return ""
                
        except Exception as e:
            logger.error(f"[DocumentParser] 读取 PDF 失败: {e}")
            return ""
        
        return text
    
    def _ocr_pdf_page(self, page) -> str:
        """对 PDF 页面进行 OCR"""
        try:
            import fitz
            
            mat = fitz.Matrix(2, 2)
            pix = page.get_pixmap(matrix=mat)
            
            img_data = pix.tobytes("png")
            
            return self._ocr_image_bytes(img_data)
            
        except Exception as e:
            logger.error(f"[DocumentParser] PDF OCR 失败: {e}")
            return ""
    
    def _read_image(self, filepath: str) -> str:
        """读取图片
        
        优先使用视觉模型生成描述，其次使用 OCR
        """
        if self._vision_model:
            return self._caption_image(filepath)
        else:
            return self._ocr_image(filepath)
    
    def _caption_image(self, filepath: str) -> str:
        """使用视觉模型生成图片描述"""
        try:
            if self._vision_model and hasattr(self._vision_model, 'generate'):
                prompt = "请详细描述这张图片的内容。"
                caption = self._vision_model.generate(prompt, image=filepath, max_tokens=200)
                return f"[图片描述] {caption}"
            else:
                return self._ocr_image(filepath)
        except Exception as e:
            logger.error(f"[DocumentParser] 图片描述生成失败: {e}")
            return self._ocr_image(filepath)
    
    def _ocr_image(self, filepath: str) -> str:
        """对图片进行 OCR"""
        try:
            with open(filepath, 'rb') as f:
                img_bytes = f.read()
            return self._ocr_image_bytes(img_bytes)
        except Exception as e:
            logger.error(f"[DocumentParser] 图片读取失败: {e}")
            return ""
    
    def _ocr_image_bytes(self, img_bytes: bytes) -> str:
        """对图片字节进行 OCR"""
        if self._ocr_engine == "tesseract":
            return self._ocr_with_tesseract(img_bytes)
        elif self._ocr_engine in ["pymupdf", "pdfplumber"]:
            return self._ocr_with_paddle(img_bytes)
        else:
            logger.warning("[DocumentParser] 无可用 OCR 引擎")
            return "[图片内容无法识别]"
    
    def _ocr_with_tesseract(self, img_bytes: bytes) -> str:
        """使用 Tesseract OCR"""
        try:
            import pytesseract
            from PIL import Image
            import io
            
            img = Image.open(io.BytesIO(img_bytes))
            text = pytesseract.image_to_string(img, lang='chi_sim+eng')
            return text.strip()
        except Exception as e:
            logger.error(f"[DocumentParser] Tesseract OCR 失败: {e}")
            return ""
    
    def _ocr_with_paddle(self, img_bytes: bytes) -> str:
        """使用 PaddleOCR"""
        try:
            from paddleocr import PaddleOCR
            
            ocr = PaddleOCR(use_angle_cls=True, lang='ch', use_gpu=False)
            
            import tempfile
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as f:
                f.write(img_bytes)
                temp_path = f.name
            
            result = ocr.ocr(temp_path, cls=True)
            
            os.unlink(temp_path)
            
            texts = []
            if result and result[0]:
                for line in result[0]:
                    texts.append(line[1][0])
            
            return "\n".join(texts)
            
        except ImportError:
            logger.warning("[DocumentParser] PaddleOCR 未安装")
            return ""
        except Exception as e:
            logger.error(f"[DocumentParser] PaddleOCR 失败: {e}")
            return ""
    
    def _chunk_content(self, content: str, chunk_size: int = 500) -> List[str]:
        """将内容分块
        
        Args:
            content: 文本内容
            chunk_size: 块大小
            
        Returns:
            List[str]: 分块列表
        """
        if not content:
            return []
        
        if len(content) <= chunk_size:
            return [content.strip()]
        
        chunks = []
        paragraphs = content.split("\n\n")
        current_chunk = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            if len(current_chunk) + len(para) + 2 > chunk_size:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                
                if len(para) > chunk_size:
                    sentences = self._split_sentences(para)
                    for sentence in sentences:
                        if len(sentence) > chunk_size:
                            for i in range(0, len(sentence), chunk_size):
                                chunks.append(sentence[i:i+chunk_size])
                        else:
                            if len(current_chunk) + len(sentence) > chunk_size:
                                chunks.append(current_chunk.strip())
                                current_chunk = sentence
                            else:
                                current_chunk += " " + sentence
                else:
                    current_chunk = para
            else:
                current_chunk += "\n\n" + para if current_chunk else para
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def _split_sentences(self, text: str) -> List[str]:
        """分割句子"""
        import re
        
        sentences = re.split(r'[。！？\n]', text)
        return [s.strip() for s in sentences if s.strip()]
    
    def _build_metadata(self, filepath: str, file_type: str, content: str) -> Dict[str, Any]:
        """构建元数据
        
        Args:
            filepath: 文件路径
            file_type: 文件类型
            content: 内容
            
        Returns:
            Dict: 元数据
        """
        import time
        
        stat = os.stat(filepath) if os.path.exists(filepath) else None
        
        return {
            "source_file": os.path.basename(filepath),
            "file_type": file_type,
            "file_size": stat.st_size if stat else 0,
            "content_length": len(content),
            "created_time": stat.st_ctime if stat else None,
            "modified_time": stat.st_mtime if stat else None,
            "ingested_at": time.time()
        }


_document_parser: Optional[DocumentParser] = None


def get_document_parser() -> DocumentParser:
    """获取文档解析器单例"""
    global _document_parser
    if _document_parser is None:
        _document_parser = DocumentParser()
        _document_parser.initialize()
    return _document_parser
